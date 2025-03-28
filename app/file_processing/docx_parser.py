from typing import Dict
import docx
from .base_parser import BaseParser

class DocxParser(BaseParser):
    def parse(self, file_path: str) -> Dict:
        try:
            doc = docx.Document(file_path)
            return {
                'text': '\n'.join([para.text for para in doc.paragraphs]),
                'metadata': {
                    'sections': [{'heading': para.style.name, 'content': para.text} 
                               for para in doc.paragraphs if para.style.name.startswith('Heading')]
                }
            }
        except Exception as e:
            self.logger.error(f"DOCX解析失败: {str(e)}")
            raise